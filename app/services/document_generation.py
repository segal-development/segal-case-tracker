"""Generate editable .docx court filings from case data.

Slice 1 (system requirement #3): the *escrito de oposición de excepciones*
for a Chilean juicio ejecutivo, defence (ejecutado) side — arts. 459/464 of
the Código de Procedimiento Civil (CPC).

The document is produced on demand and streamed back to the lawyer, who edits
it and uploads it to PJUD manually. Nothing is persisted here: no Document row,
no storage write. The legal CONTENT is intentionally in Spanish (it is a
Chilean court filing); every identifier, comment and docstring stays in English.

Fields that are missing/empty are rendered as bracketed, editable placeholders
so the lawyer can complete them by hand.
"""

from __future__ import annotations

import io
from dataclasses import dataclass
from datetime import date
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from app.services.lawyer_roster import _clean_nombre

# Título types whose executive action prescribes in ONE year under the Ley
# sobre Letras de Cambio y Pagarés (DFL N°707, art. 98). Everything else falls
# back to the generic three-year executive-action prescription (art. 2515 CC).
_ONE_YEAR_TITULO_TIPOS = frozenset({"pagare", "letra", "cheque"})

_ONE_YEAR_PLAZO_PHRASE = (
    "el plazo de prescripción de un año que establece el artículo 98 del "
    "DFL N°707 (Ley sobre Letras de Cambio y Pagarés)"
)
_THREE_YEAR_PLAZO_PHRASE = (
    "el plazo de prescripción de tres años de la acción ejecutiva que "
    "establece el artículo 2515 del Código Civil"
)

# Spanish month names — kept explicit so date formatting never depends on the
# process locale (which is not guaranteed to carry an es_CL/es_ES catalogue).
_SPANISH_MONTHS = (
    "enero", "febrero", "marzo", "abril", "mayo", "junio",
    "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre",
)


@dataclass
class Party:
    """A resolved litigant (ejecutante or ejecutado) for the filing."""

    nombre: str
    rut: str
    persona_type: str


def _format_spanish_date(value: Optional[date]) -> Optional[str]:
    """Format a date as ``"5 de marzo de 2024"`` without relying on locale.

    Returns ``None`` when there is no date, so callers can substitute their own
    bracketed placeholder.
    """
    if value is None:
        return None
    return f"{value.day} de {_SPANISH_MONTHS[value.month - 1]} de {value.year}"


def resolve_parties(litigantes) -> tuple[Optional[Party], Optional[Party]]:
    """Resolve (ejecutante, ejecutado) from a case's litigantes.

    EJECUTANTE is the litigante coded ``DTE.`` (demandante); EJECUTADO is the
    one coded ``DDO.`` (demandado). Names are cleaned of any trailing
    parenthetical (e.g. a "(APODERADO)" suffix) reusing the roster helper.
    Either side may be ``None`` when absent from the litigantes list.
    """
    ejecutante: Optional[Party] = None
    ejecutado: Optional[Party] = None

    for lit in litigantes or []:
        participante = (getattr(lit, "participante", "") or "").strip()
        party = Party(
            nombre=_clean_nombre(getattr(lit, "nombre", "") or ""),
            rut=(getattr(lit, "rut", "") or "").strip(),
            persona_type=(getattr(lit, "persona_type", "") or "").strip(),
        )
        if participante == "DTE." and ejecutante is None:
            ejecutante = party
        elif participante == "DDO." and ejecutado is None:
            ejecutado = party

    return ejecutante, ejecutado


def _add_paragraph(doc: Document, *, justify: bool = True):
    """Add an empty paragraph, justified by default, and return it."""
    para = doc.add_paragraph()
    if justify:
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return para


def _add_run(para, text: str, *, bold: bool = False) -> None:
    run = para.add_run(text)
    run.bold = bold


def build_escrito_oposicion(
    case,
    litigantes,
    court_name: Optional[str],
    acting_lawyer_name: Optional[str],
    acting_lawyer_rut: Optional[str],
) -> bytes:
    """Build the *escrito de oposición de excepciones* and return docx bytes.

    All inputs may be missing/empty; each is rendered as a bracketed, editable
    placeholder in that case. No system suggestion/disclaimer text is embedded
    in the filing — that is a UI concern, not part of the court document.
    """
    ejecutante, ejecutado = resolve_parties(litigantes)

    # Prefer the DTE./DDO. litigantes (they carry a RUT), but fall back to the
    # case's free-text plaintiff/defendant. Only ~17% of cases have DTE./DDO.
    # litigantes while ~98% have plaintiff/defendant, so this fills the
    # caratulado (name-only; RUT stays a placeholder) far more often.
    case_plaintiff = (getattr(case, "plaintiff", None) or "").strip()
    case_defendant = (getattr(case, "defendant", None) or "").strip()

    ejecutante_nombre = (
        (ejecutante.nombre if ejecutante else "") or case_plaintiff or "[EJECUTANTE]"
    )
    ejecutado_nombre = (
        (ejecutado.nombre if ejecutado else "") or case_defendant or "[EJECUTADO]"
    )
    ejecutado_rut_txt = (
        f"RUT {ejecutado.rut}" if (ejecutado and ejecutado.rut) else "RUT [___]"
    )

    lawyer_name = acting_lawyer_name or "[NOMBRE ABOGADO]"
    lawyer_rut = acting_lawyer_rut or "[RUT ABOGADO]"

    rol = getattr(case, "rol", None) or "[ROL]"

    doc = Document()

    # Readable body font + justified default via the Normal style.
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    # ── SUMA (bold header) ────────────────────────────────────────────────
    suma = _add_paragraph(doc)
    _add_run(
        suma,
        "EN LO PRINCIPAL: OPONE EXCEPCIONES A LA EJECUCIÓN; PRIMER OTROSÍ: "
        "ACOMPAÑA DOCUMENTOS; SEGUNDO OTROSÍ: PATROCINIO Y PODER.",
        bold=True,
    )

    _add_paragraph(doc, justify=False)  # blank line

    # ── S.J.L. addressee ──────────────────────────────────────────────────
    court_suffix = f" de {court_name}" if court_name else " (___)"
    sjl = _add_paragraph(doc, justify=False)
    _add_run(sjl, "S.J.L.", bold=True)
    _add_run(sjl, f" en lo Civil{court_suffix}")

    _add_paragraph(doc, justify=False)  # blank line

    # ── Intro / comparecencia ─────────────────────────────────────────────
    intro = _add_paragraph(doc)
    _add_run(
        intro,
        f"{lawyer_name}, abogado, RUT {lawyer_rut}, en representación "
        f"convencional de {ejecutado_nombre}, {ejecutado_rut_txt}, parte "
        f"ejecutada en los autos ejecutivos caratulados "
        f'"{ejecutante_nombre} con {ejecutado_nombre}", Rol {rol}, a US. '
        f"respetuosamente digo:",
    )

    _add_paragraph(doc, justify=False)  # blank line

    # ── EN LO PRINCIPAL body ──────────────────────────────────────────────
    body = _add_paragraph(doc)
    _add_run(
        body,
        "Que, encontrándome dentro del plazo legal de ocho días hábiles que "
        "establece el artículo 459 del Código de Procedimiento Civil, vengo en "
        "oponer a la ejecución las siguientes excepciones, de conformidad a lo "
        "dispuesto en el artículo 464 del mismo Código:",
    )

    _add_paragraph(doc, justify=False)  # blank line

    # ── Excepciones (numbered) ────────────────────────────────────────────
    number = 1
    if getattr(case, "prescripcion_cumplida", False):
        titulo_tipo = (getattr(case, "titulo_tipo", None) or "").strip().lower()
        plazo_phrase = (
            _ONE_YEAR_PLAZO_PHRASE
            if titulo_tipo in _ONE_YEAR_TITULO_TIPOS
            else _THREE_YEAR_PLAZO_PHRASE
        )
        titulo_fecha_txt = (
            _format_spanish_date(getattr(case, "titulo_fecha", None))
            or "[FECHA DEL TÍTULO]"
        )

        presc = _add_paragraph(doc)
        _add_run(
            presc,
            f"{number}. EXCEPCIÓN DE PRESCRIPCIÓN DE LA ACCIÓN EJECUTIVA "
            "(artículo 464 N°17 del Código de Procedimiento Civil). ",
            bold=True,
        )
        _add_run(
            presc,
            "Opongo la excepción de prescripción, por cuanto a la fecha de "
            "notificación de la demanda se encontraba cumplido "
            f"{plazo_phrase}, contado desde {titulo_fecha_txt}. En "
            "consecuencia, la acción ejecutiva se encuentra prescrita y así "
            "solicito se declare.",
        )
        number += 1

    placeholder = _add_paragraph(doc)
    _add_run(
        placeholder,
        f"{number}. [INDICAR AQUÍ LAS DEMÁS EXCEPCIONES QUE CORRESPONDAN "
        "conforme al artículo 464 del Código de Procedimiento Civil — por "
        "ejemplo, N°7 (falta de alguno de los requisitos o condiciones "
        "establecidos por las leyes para que el título tenga fuerza "
        "ejecutiva), acompañando los fundamentos de hecho y de derecho.]",
    )

    _add_paragraph(doc, justify=False)  # blank line

    # ── POR TANTO ─────────────────────────────────────────────────────────
    por_tanto = _add_paragraph(doc)
    _add_run(por_tanto, "POR TANTO,", bold=True)
    _add_run(
        por_tanto,
        " y de conformidad a lo dispuesto en los artículos 459, 464 y demás "
        "pertinentes del Código de Procedimiento Civil,",
    )

    ruego = _add_paragraph(doc)
    _add_run(ruego, "RUEGO A US.:", bold=True)
    _add_run(
        ruego,
        " Tener por opuestas las excepciones precedentes dentro de plazo, "
        "acogerlas a tramitación y, en definitiva, rechazar la ejecución en "
        "todas sus partes, con expresa condena en costas.",
    )

    _add_paragraph(doc, justify=False)  # blank line

    # ── PRIMER OTROSÍ ─────────────────────────────────────────────────────
    primer = _add_paragraph(doc)
    _add_run(primer, "PRIMER OTROSÍ:", bold=True)
    _add_run(
        primer,
        " Ruego a US. tener por acompañados, con citación, los siguientes "
        "documentos en que fundo las excepciones opuestas: [INDICAR "
        "DOCUMENTOS].",
    )

    # ── SEGUNDO OTROSÍ ────────────────────────────────────────────────────
    segundo = _add_paragraph(doc)
    _add_run(segundo, "SEGUNDO OTROSÍ:", bold=True)
    _add_run(
        segundo,
        " Ruego a US. tener presente que designo como abogado patrocinante y "
        f"confiero poder a {lawyer_name}, RUT {lawyer_rut}, con domicilio en "
        "[DOMICILIO], quien firma en señal de aceptación.",
    )

    # ── Signature block ───────────────────────────────────────────────────
    _add_paragraph(doc, justify=False)  # blank line
    _add_paragraph(doc, justify=False)  # blank line

    sig_line = _add_paragraph(doc, justify=False)
    sig_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(sig_line, "_______________________________")

    sig_name = _add_paragraph(doc, justify=False)
    sig_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(sig_name, lawyer_name)

    sig_rut = _add_paragraph(doc, justify=False)
    sig_rut.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(sig_rut, f"RUT {lawyer_rut}")

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
