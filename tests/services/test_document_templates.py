"""Tests for the template-based document generator (document_templates.py).

Deterministic and offline: they render the REAL committed .docx template with a
fake case + litigantes and assert the derived variables fill in, unknown facts
become «INDICAR …» prompts, and no raw {{ placeholder }} survives.
"""
import io
import re
from datetime import date, datetime
from types import SimpleNamespace

import pytest
from docx import Document

from app.services.document_templates import (
    TEMPLATE_REGISTRY,
    _derive_sjl,
    is_template_document,
    render_template_document,
)


def _lit(participante, nombre, rut="", persona_type="natural"):
    return SimpleNamespace(participante=participante, nombre=nombre, rut=rut, persona_type=persona_type)


def _render(case, litigantes, court_name, **lawyer):
    data = render_template_document(
        document_type="abandono_3anios",
        case=case,
        litigantes=litigantes,
        court_name=court_name,
        acting_lawyer_name=lawyer.get("name"),
        acting_lawyer_rut=lawyer.get("rut"),
        acting_lawyer_email=lawyer.get("email"),
    )
    text = "\n".join(p.text for p in Document(io.BytesIO(data)).paragraphs)
    return data, text


def test_abandono_renders_with_full_case_data():
    case = SimpleNamespace(
        rol="C-12111-2025", plaintiff="BANCO FALABELLA",
        defendant="ZULEMA GARRIDO MORENO", last_movement_at=datetime(2025, 10, 1),
    )
    lits = [
        _lit("DTE.", "BANCO FALABELLA S.A.", "96509660-4"),
        _lit("DDO.", "ZULEMA VERÓNICA GARRIDO MORENO", "18995444-1"),
    ]
    data, text = _render(
        case, lits, "10º Juzgado Civil de Santiago",
        name="FERNANDA ARROYO DIAZ", rut="18914338-9", email="farroyo@segal.cl",
    )
    assert len(data) > 5000  # a real .docx
    # Header + parties filled from case/litigantes.
    assert "C-12111-2025" in text
    assert "10º Juzgado Civil de Santiago" in text
    assert "ZULEMA VERÓNICA GARRIDO MORENO" in text
    assert "S.J.L. Civil de Santiago (10°)" in text  # derived
    # RUTs formatted with thousands separators.
    assert "18.995.444-1" in text  # ejecutado
    assert "18.914.338-9" in text  # abogado
    assert "farroyo@segal.cl" in text
    # Facts the system can't derive are visible prompts.
    assert "«INDICAR PROFESIÓN DEL EJECUTADO»" in text
    assert "«INDICAR MONTO DEMANDADO»" in text
    assert "aprox. 1 de octubre de 2025" in text  # última-gestión hint
    # Nothing left unrendered.
    assert re.search(r"\{\{.*?\}\}", text) is None


def test_missing_data_falls_back_to_prompts_without_crashing():
    case = SimpleNamespace(rol=None, plaintiff=None, defendant=None, last_movement_at=None)
    data, text = _render(case, [], None)
    assert len(data) > 5000
    assert "«INDICAR ROL»" in text
    assert "«INDICAR EJECUTADO»" in text
    assert "«INDICAR ABOGADO PATROCINANTE»" in text
    assert "«INDICAR FECHA ÚLTIMA GESTIÓN ÚTIL»" in text  # no hint when no movement
    assert re.search(r"\{\{.*?\}\}", text) is None


def test_registry_maps_document_to_its_recommendation():
    assert is_template_document("abandono_3anios") is True
    assert is_template_document("escrito_oposicion") is False
    spec = TEMPLATE_REGISTRY["abandono_3anios"]
    assert spec.recommendation_code == "solicitar_abandono"
    assert spec.template_filename == "abandono_3anios.docx"


@pytest.mark.parametrize(
    "court,expected",
    [
        ("10º Juzgado Civil de Santiago", "S.J.L. Civil de Santiago (10°)"),
        ("1º Juzgado Civil de San Miguel", "S.J.L. Civil de San Miguel (1°)"),
        (None, "«INDICAR TRIBUNAL»"),
    ],
)
def test_derive_sjl(court, expected):
    assert _derive_sjl(court) == expected


def _render_type(document_type, case, litigantes, court_name, **lawyer):
    data = render_template_document(
        document_type=document_type,
        case=case,
        litigantes=litigantes,
        court_name=court_name,
        acting_lawyer_name=lawyer.get("name"),
        acting_lawyer_rut=lawyer.get("rut"),
        acting_lawyer_email=lawyer.get("email"),
    )
    text = "\n".join(p.text for p in Document(io.BytesIO(data)).paragraphs)
    return data, text


# The .docx templates were built from real filings; these litmus values from the
# source cases must NEVER survive into a rendered document (data-hygiene guard).
_SOURCE_CASE_LEAKS = [
    "NELSON", "1,85%", "169.636", "5.572.062", "farroyo", "bhervas",
    "31-08-2022", "ITAU", "SANDRA", "BENJAMIN", "RICARDO", "15º Juzgado",
    "C-18599", "C-20862",
]


def test_prescripcion_renders_and_prompts_non_derivable_facts():
    case = SimpleNamespace(
        rol="C-777-2026", plaintiff="CAJA DE COMPENSACIÓN",
        defendant="MARÍA SOTO", last_movement_at=None,
    )
    lits = [
        _lit("DTE.", "CAJA DE COMPENSACIÓN LOS HÉROES", "81111000-0"),
        _lit("DDO.", "MARÍA SOTO PÉREZ", "17222333-8"),
    ]
    data, text = _render_type(
        "prescripcion_cuotas", case, lits, "11º Juzgado Civil de Santiago",
        name="ABOGADA TEST", rut="16111222-4", email="abg@segal.cl",
    )
    assert len(data) > 5000
    # Derivable header/parties/patrocinio fill in.
    assert "C-777-2026" in text
    assert "S.J.L. Civil de Santiago (11°)" in text
    assert "MARÍA SOTO PÉREZ" in text
    assert "17.222.333-8" in text  # ejecutado rut formatted
    assert "81.111.000-0" in text  # ejecutante rut formatted
    assert "abg@segal.cl" in text  # notificación email
    # Pagaré facts (none derivable from PJUD) become visible prompts.
    assert "«INDICAR MONTO DEMANDADO (CAPITAL)»" in text
    assert "«INDICAR TASA DE INTERÉS»" in text
    assert "«INDICAR N° DE CUOTAS»" in text
    assert "«INDICAR SALDO ADEUDADO»" in text
    assert "«INDICAR REPRESENTANTE Y DOMICILIO DEL EJECUTANTE»" in text
    # No raw placeholder and no figure from the source filing.
    assert re.search(r"\{\{.*?\}\}", text) is None
    for leak in _SOURCE_CASE_LEAKS:
        assert leak not in text, f"source-case value leaked: {leak}"


def test_objeta_remate_renders_all_derivable():
    case = SimpleNamespace(
        rol="C-555-2026", plaintiff="BANCO DE PRUEBA",
        defendant="PEDRO DÍAZ", last_movement_at=None,
    )
    lits = [
        _lit("DTE.", "BANCO DE PRUEBA S.A.", "97000000-3"),
        _lit("DDO.", "PEDRO DÍAZ ROJAS", "16111222-4"),
    ]
    data, text = _render_type(
        "objeta_remate", case, lits, "5º Juzgado Civil de Santiago",
        name="ABOGADO TEST", rut="17222333-8", email="abg@segal.cl",
    )
    assert len(data) > 5000
    assert "C-555-2026" in text
    assert "S.J.L. Civil de Santiago (5°)" in text
    assert "BANCO DE PRUEBA S.A." in text  # ejecutante in header + body clause
    assert "ABOGADO TEST" in text
    # Header + fixed-clause body: no prompts, no leftover source data.
    assert re.search(r"\{\{.*?\}\}", text) is None
    for leak in _SOURCE_CASE_LEAKS:
        assert leak not in text, f"source-case value leaked: {leak}"


def test_new_templates_are_registered():
    for dt, rec in (
        ("prescripcion_cuotas", "oponer_excepciones"),
        ("objeta_remate", "objetar_remate"),
    ):
        assert is_template_document(dt) is True
        assert TEMPLATE_REGISTRY[dt].recommendation_code == rec
