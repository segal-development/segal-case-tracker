"""Tests for the escrito-de-oposición generator (system requirement #3, slice 1).

Covers the pure service (party resolution + non-empty docx bytes) and the
streaming endpoint (auth/scope + docx attachment response).
"""

from datetime import date, datetime, timedelta
from types import SimpleNamespace

import pytest

from app.core.security import create_access_token
from app.models.case import Case
from app.models.case_litigante import CaseLitigante
from app.models.court import Court
from app.models.lawyer import Lawyer
from app.services.document_generation import (
    build_escrito_oposicion,
    resolve_firm_side,
    resolve_parties,
)

DOCX_MEDIA_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
AUDITOR_RUT = "77777777-7"


def _lit(participante, rut, nombre, persona_type="NATURAL"):
    return SimpleNamespace(
        participante=participante,
        rut=rut,
        nombre=nombre,
        persona_type=persona_type,
    )


class TestResolveParties:
    def test_picks_dte_as_ejecutante_and_ddo_as_ejecutado(self):
        litigantes = [
            _lit("DDO.", "12345678-9", "Juan Pérez (DEUDOR)"),
            _lit("DTE.", "60000000-0", "Banco Ejemplo S.A."),
            _lit("AB.DTE", "11111111-1", "Abogado Contrario"),
        ]
        ejecutante, ejecutado = resolve_parties(litigantes)
        assert ejecutante is not None and ejecutado is not None
        assert ejecutante.nombre == "Banco Ejemplo S.A."
        assert ejecutado.nombre == "Juan Pérez"  # trailing parenthetical stripped
        assert ejecutado.rut == "12345678-9"

    def test_missing_sides_return_none(self):
        ejecutante, ejecutado = resolve_parties([])
        assert ejecutante is None and ejecutado is None


class TestResolveFirmSide:
    """The escrito must never ASSERT a representation it cannot verify.

    deadline_engine._firm_side collapses "confirmed demandado" and "no idea"
    into "demandado". For a court filing that is not acceptable — an unmatched
    lawyer must surface as "desconocido", not as a silent claim of defence.
    """

    def test_ab_ddo_match_is_demandado(self):
        lits = [_lit("AB.DDO", "9999999-9", "Abogada Defensora"),
                _lit("DDO.", "12345678-9", "Juan Pérez")]
        assert resolve_firm_side(lits, "9999999-9") == "demandado"

    def test_ap_ddo_match_is_demandado(self):
        assert resolve_firm_side([_lit("AP.DDO", "9999999-9", "X")], "9999999-9") == "demandado"

    def test_ab_dte_match_is_demandante(self):
        lits = [_lit("AB.DTE", "9999999-9", "Abogado del Banco")]
        assert resolve_firm_side(lits, "9999999-9") == "demandante"

    def test_no_litigantes_is_desconocido_not_demandado(self):
        assert resolve_firm_side([], "9999999-9") == "desconocido"

    def test_lawyer_not_among_litigantes_is_desconocido(self):
        lits = [_lit("AB.DDO", "1111111-1", "Otro Abogado")]
        assert resolve_firm_side(lits, "9999999-9") == "desconocido"

    def test_missing_acting_rut_is_desconocido(self):
        assert resolve_firm_side([_lit("AB.DDO", "9999999-9", "X")], None) == "desconocido"

    def test_rut_matching_ignores_dots(self):
        lits = [_lit("AB.DDO", "9.999.999-9", "Abogada")]
        assert resolve_firm_side(lits, "9999999-9") == "demandado"


class TestBuildEscrito:
    def _case(self, **overrides):
        defaults = dict(
            rol="C-1234-2025",
            titulo_tipo=None,
            titulo_fecha=None,
            prescripcion_cumplida=False,
        )
        defaults.update(overrides)
        return SimpleNamespace(**defaults)

    def test_returns_non_empty_docx_bytes(self):
        data = build_escrito_oposicion(
            case=self._case(),
            litigantes=[
                _lit("DTE.", "60000000-0", "Banco Ejemplo S.A."),
                _lit("DDO.", "12345678-9", "Juan Pérez"),
            ],
            court_name="Santiago",
            acting_lawyer_name="María González",
            acting_lawyer_rut="9999999-9",
        )
        assert isinstance(data, bytes)
        assert len(data) > 0
        assert data[:2] == b"PK"  # .docx is a zip container

    def test_prescripcion_pagare_uses_one_year_basis(self):
        from docx import Document
        import io

        data = build_escrito_oposicion(
            case=self._case(
                titulo_tipo="pagare",
                titulo_fecha=date(2020, 3, 5),
                prescripcion_cumplida=True,
            ),
            litigantes=[],
            court_name=None,
            acting_lawyer_name=None,
            acting_lawyer_rut=None,
        )
        text = "\n".join(p.text for p in Document(io.BytesIO(data)).paragraphs)
        assert "artículo 98 del DFL N°707" in text
        assert "contado desde el vencimiento del documento" in text
        assert "5 de marzo de 2020" in text
        assert "[EJECUTADO]" in text  # placeholder when litigante absent

    def test_prescripcion_cheque_uses_art_34_from_protesto(self):
        from docx import Document
        import io

        # Per the firm's reference doc: cheque prescribes in 1 year counted from
        # the PROTESTO (art. 34 DFL 707) — not art. 98, and not from the vencimiento.
        data = build_escrito_oposicion(
            case=self._case(
                titulo_tipo="cheque",
                titulo_fecha=date(2021, 11, 8),
                prescripcion_cumplida=True,
            ),
            litigantes=[],
            court_name=None,
            acting_lawyer_name=None,
            acting_lawyer_rut=None,
        )
        text = "\n".join(p.text for p in Document(io.BytesIO(data)).paragraphs)
        assert "artículo 34 del DFL N°707" in text
        assert "Cuentas Corrientes Bancarias y Cheques" in text
        assert "contado desde la fecha del protesto" in text
        assert "8 de noviembre de 2021" in text
        assert "artículo 98" not in text  # must NOT reuse the cambiaria citation

    def test_prescripcion_cheque_without_date_uses_protesto_placeholder(self):
        from docx import Document
        import io

        data = build_escrito_oposicion(
            case=self._case(titulo_tipo="cheque", prescripcion_cumplida=True),
            litigantes=[],
            court_name=None,
            acting_lawyer_name=None,
            acting_lawyer_rut=None,
        )
        text = "\n".join(p.text for p in Document(io.BytesIO(data)).paragraphs)
        assert "[FECHA DEL PROTESTO]" in text

    def test_unconfirmed_representation_is_flagged_not_asserted(self):
        from docx import Document
        import io

        # No AB.DDO litigante matches the acting lawyer → the system cannot know
        # whom it represents. It must SAY so instead of silently naming a client.
        data = build_escrito_oposicion(
            case=self._case(plaintiff="Barrera", defendant="Promotora CMR Falabella S.A."),
            litigantes=[],
            court_name=None,
            acting_lawyer_name="Abogada X",
            acting_lawyer_rut="9999999-9",
            firm_side="desconocido",
        )
        text = "\n".join(p.text for p in Document(io.BytesIO(data)).paragraphs)
        assert "[CONFIRMAR REPRESENTACIÓN" in text

    def test_confirmed_defence_has_no_confirmation_marker(self):
        from docx import Document
        import io

        data = build_escrito_oposicion(
            case=self._case(),
            litigantes=[_lit("DDO.", "12345678-9", "Juan Pérez"),
                        _lit("AB.DDO", "9999999-9", "Abogada X")],
            court_name=None,
            acting_lawyer_name="Abogada X",
            acting_lawyer_rut="9999999-9",
            firm_side="demandado",
        )
        text = "\n".join(p.text for p in Document(io.BytesIO(data)).paragraphs)
        assert "[CONFIRMAR REPRESENTACIÓN" not in text
        assert "Juan Pérez" in text

    def test_prescripcion_generic_titulo_uses_three_year_2515(self):
        from docx import Document
        import io

        data = build_escrito_oposicion(
            case=self._case(titulo_tipo="escritura_publica", prescripcion_cumplida=True),
            litigantes=[],
            court_name=None,
            acting_lawyer_name=None,
            acting_lawyer_rut=None,
        )
        text = "\n".join(p.text for p in Document(io.BytesIO(data)).paragraphs)
        assert "artículo 2515 del Código Civil" in text
        assert "que la obligación se hizo exigible" in text
        assert "[FECHA DEL TÍTULO]" in text

    def test_falls_back_to_case_plaintiff_defendant_without_litigantes(self):
        from docx import Document
        import io

        # ~83% of cases have no DTE./DDO. litigantes but ~98% carry the
        # free-text plaintiff/defendant — the caratulado must use those.
        data = build_escrito_oposicion(
            case=self._case(plaintiff="Promotora CMR Falabella S.A.", defendant="Sáez"),
            litigantes=[],
            court_name=None,
            acting_lawyer_name=None,
            acting_lawyer_rut=None,
        )
        text = "\n".join(p.text for p in Document(io.BytesIO(data)).paragraphs)
        assert '"Promotora CMR Falabella S.A. con Sáez"' in text
        assert "[EJECUTANTE]" not in text
        assert "[EJECUTADO]" not in text
        assert "RUT [___]" in text  # RUT still a placeholder (not in the free-text fields)


@pytest.fixture
def auditor(db):
    obj = Lawyer(rut=AUDITOR_RUT, name="Auditor", role="auditor")
    db.add(obj); db.commit(); db.refresh(obj)
    return obj


@pytest.fixture
def court(db):
    obj = Court(code="T1-DOCGEN", name="Juzgado Docgen", region="RM", type="civil")
    db.add(obj); db.commit(); db.refresh(obj)
    return obj


@pytest.fixture
def case(db, auditor, court):
    obj = Case(
        lawyer_id=auditor.id,
        court_id=court.id,
        rol="C-4502-2025",
        status="active",
        competencia="civil",
        plaintiff="Banco Ejemplo S.A.",
        defendant="Juan Pérez",
        created_at=datetime.utcnow(),
        updated_at=datetime.utcnow(),
    )
    db.add(obj); db.commit(); db.refresh(obj)
    for i, (part, rut, nombre) in enumerate([
        ("DTE.", "60000000-0", "Banco Ejemplo S.A."),
        ("DDO.", "12345678-9", "Juan Pérez"),
    ]):
        db.add(CaseLitigante(
            case_id=obj.id, participante=part, rut=rut,
            persona_type="NATURAL", nombre=nombre, natural_key=f"{obj.id}-{i}",
        ))
    db.commit()
    return obj


def _headers(rut: str) -> dict:
    tok = create_access_token({"sub": rut}, expires_delta=timedelta(minutes=30))
    return {"Authorization": f"Bearer {tok}"}


class TestGenerateEndpoint:
    def test_generates_docx_attachment(self, client, case, auditor):
        resp = client.post(
            f"/api/v1/cases/{case.id}/documents/generate",
            json={"document_type": "escrito_oposicion"},
            headers=_headers(AUDITOR_RUT),
        )
        assert resp.status_code == 200
        assert resp.headers["content-type"] == DOCX_MEDIA_TYPE
        assert "attachment" in resp.headers["content-disposition"]
        assert "oposicion_excepciones_C_4502_2025.docx" in resp.headers["content-disposition"]
        assert resp.content[:2] == b"PK"

    def test_unknown_document_type_rejected(self, client, case, auditor):
        resp = client.post(
            f"/api/v1/cases/{case.id}/documents/generate",
            json={"document_type": "nope"},
            headers=_headers(AUDITOR_RUT),
        )
        assert resp.status_code == 422

    def test_refuses_when_the_firm_represents_the_ejecutante(self, db, client, case, auditor):
        # The escrito de oposición is filed BY the ejecutado. If our lawyer is the
        # creditor's abogado (AB.DTE) the document is the wrong instrument — the
        # endpoint must refuse rather than emit a filing that contradicts itself.
        db.add(CaseLitigante(
            case_id=case.id, participante="AB.DTE", rut=AUDITOR_RUT,
            persona_type="NATURAL", nombre="Auditor", natural_key=f"{case.id}-abdte",
        ))
        db.commit()
        resp = client.post(
            f"/api/v1/cases/{case.id}/documents/generate",
            json={"document_type": "escrito_oposicion"},
            headers=_headers(AUDITOR_RUT),
        )
        assert resp.status_code == 409
        assert "ejecutante" in resp.json()["detail"].lower()

    def test_missing_case_returns_404(self, client, auditor):
        resp = client.post(
            "/api/v1/cases/999999/documents/generate",
            json={"document_type": "escrito_oposicion"},
            headers=_headers(AUDITOR_RUT),
        )
        assert resp.status_code == 404
